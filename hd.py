import docx
import os
import re
from collections import defaultdict

from docx import Document

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

filename = 'D:\Myfolder\Study\py\ChatterBot\SRS_SMART.docx'
doc = Document(filename)

# document = Document(r'D:\Myfolder\Study\py\ChatterBot\SRS_SMART.docx')
sections = doc.sections
for section in sections: print(section.start_type)
# words = document.xpath('//w:r', namespaces=document.nsmap)
# document = opendocx(r'D:\Myfolder\Study\py\ChatterBot\SRS_SMART.docx')

len(doc.paragraphs)
len(doc.tables)

from docx.enum.style import WD_STYLE_TYPE

styles = doc.styles

paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
for style in paragraph_styles: print(style.name)


def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


# data=dict()
data = {}

t = doc.element.body.iterchildren()
t = t.next()


def table_print(block):
    table = block
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                # print(paragraph.text,'  ', end=" ")
                print(paragraph.text, '  ')
                # y.write(paragraph.text)
                # y.write('  ')
        print("\n")


for block in iter_block_items(doc):
    print('found one', block.__class__.__name__)
    if isinstance(block, Paragraph):
        data[block.__class__.__name__] = block.text
        # data[block.__class__.__name__] = data[block.__class__.__name__] + block.text
    else:
        data[block.__class__.__name__] = table_print(block)


def docx_replace_regex(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)


# for p in doc.paragraphs:
# print len(p.text)

# regex1 = re.compile(r"your regex")
# replace1 = r"your replace string"

# for word, replacement in dictionary.items():
# word_re=re.compile(word)
# docx_replace_regex(doc, word_re , replacement)

# docx_replace_regex(doc, regex1 , replace1)
# doc.save('result1.docx')

def makedict():
    with open('textfile.txt') as f:
        words = [i.strip().lower() for i in f.read().split()]
        return dict(zip(words[:-1], words[1:]))


def strip(value):
    """ Strip the input value if it is a string and returns None
        if it had only whitespaces """
    if isinstance(value, basestring):
        value = value.strip()
        if len(value) == 0:
            return None
    return value
